import csv
from docx import Document
import analysis_tools as at


# Class for a datafile
class DataFile:
    def __init__(self, filename, title):
        self.filename = filename
        self.title = title
        self.description = None
        self.process = []
        self.data = []
        self.header = []

    def set_title(self, x):
        self.title = x

    def get_title(self):
        return self.title

    def set_description(self, x):
        self.description = x

    def get_description(self):
        return self.description

    def set_process(self, x):
        self.process = x

    def get_process(self):
        return self.process

    def set_data(self, x):
        self.data = x

    def get_data(self):
        return self.data

    def set_header(self, x):
        self.header = x

    def get_header(self):
        return self.header

    def get_col_names(self, index):
        return self.header[index]

    def get_col(self, col):
        """ Gets a list holding all the elements of a specific column """
        newList = []
        # Goes through each list within the lists of data
        for row in self.data:
            # Appends an element from each row at index of passed in col
            newList.append(row[col])
        return newList

    def load(self):
        """ Reads and opens a CSV file, assigns headers to instance of datafile """
        data = []
        with open(self.filename) as csvfile:
            reader = csv.reader(csvfile)
            # Appends each row of the csv file into self.data
            for row in reader:
                data.append(row)
            # Puts the first row of data into self.header and dels it from data
            self.header = data[0]
            del(data[0])
            # Puts data into data (instance of datafile)
            self.data = data


class ReportGenerator:
    def __init__(self, datafile):
        self.datafile = datafile

    def generate(self, output_name):
        """ Generates a document with the information on sales, etc... """
        # Creates instance of a document using python-docx
        document = Document()
        # Adds heading to document of the datafile title
        document.add_heading(self.datafile.get_title(), 0)
        # Adds text to document of datafile description
        document.add_paragraph(self.datafile.get_description())

        # Iterates through data needed (Scores - Other Sales)
        #for i in range(6, (len(self.datafile.header) - 1)):
        for i in self.datafile.get_process():
            # Adds heading to document from datafile header at looped index
            document.add_heading(self.datafile.header[i])
            # Adds table to document w/ 2 columns
            table = document.add_table(rows=1, cols=2)

            # Runs functions on column of data using get_col function
            new_list = at.MyMath.cleanup(self.datafile.get_col(i))
            my_mean = at.MyMath.mean(new_list)
            my_stddev = at.MyMath.stddev(new_list)
            my_median = at.MyMath.median(new_list)
            my_min = at.MyMath.min(new_list)
            my_max = at.MyMath.max(new_list)

            # Adds row to table and returns cells to produce label on
            # left side of table and the value from functions on right
            row_cells = table.add_row().cells
            row_cells[0].text = 'Mean:'
            row_cells[1].text = str(my_mean)

            row_cells = table.add_row().cells
            row_cells[0].text = 'Standard Deviation:'
            row_cells[1].text = str(my_stddev)

            row_cells = table.add_row().cells
            row_cells[0].text = 'Median:'
            row_cells[1].text = str(my_median)

            row_cells = table.add_row().cells
            row_cells[0].text = 'Minimum:'
            row_cells[1].text = str(my_min)

            row_cells = table.add_row().cells
            row_cells[0].text = 'Maximum:'
            row_cells[1].text = str(my_max)

        # Saves the document as the passed in output_name
        document.save(output_name)

